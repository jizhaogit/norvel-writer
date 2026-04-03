from __future__ import annotations
from pathlib import Path
from norvel_writer.ingestion.base import BaseIngestor, IngestedDocument


# Map python-docx paragraph style names → Markdown heading prefix
_HEADING_STYLES: dict[str, str] = {
    "Title":     "#",
    "Subtitle":  "##",
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Heading 5": "#####",
    "Heading 6": "######",
}

# Style names whose paragraphs are unordered list items
_BULLET_STYLES: set[str] = {
    "List Bullet", "List Bullet 2", "List Bullet 3",
    "List Paragraph",
}

# Style names whose paragraphs are ordered list items
_NUMBER_STYLES: set[str] = {
    "List Number", "List Number 2", "List Number 3",
    "List Continue",
}


def _runs_to_markdown(para) -> str:
    """
    Convert python-docx paragraph runs to Markdown inline markup.
    Handles bold, italic, and bold-italic combinations.
    Runs with no text are skipped.
    """
    parts: list[str] = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        parts.append(t)
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """
    Convert a python-docx Table to a GitHub-Flavoured Markdown table.
    The first row is treated as the header row.
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _docx_to_markdown(doc) -> str:
    """
    Walk the document body in DOM order so paragraphs and tables are
    interleaved correctly, then convert each block to Markdown.
    """
    blocks: list[str] = []

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "p":
            # Re-wrap as a python-docx Paragraph to access style / runs
            from docx.text.paragraph import Paragraph as _P
            para = _P(child, doc)
            style_name = para.style.name if para.style else ""
            raw_text = para.text.strip()
            if not raw_text:
                continue

            if style_name in _HEADING_STYLES:
                blocks.append(f"{_HEADING_STYLES[style_name]} {raw_text}")
            elif style_name in _BULLET_STYLES:
                blocks.append(f"- {raw_text}")
            elif style_name in _NUMBER_STYLES:
                blocks.append(f"1. {raw_text}")
            else:
                # Regular paragraph — preserve inline formatting
                inline = _runs_to_markdown(para)
                if inline.strip():
                    blocks.append(inline)

        elif local == "tbl":
            from docx.table import Table as _T
            table = _T(child, doc)
            md_table = _table_to_markdown(table)
            if md_table:
                blocks.append(md_table)

    return "\n\n".join(blocks)


class DocxIngestor(BaseIngestor):
    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def ingest(self, path: Path) -> IngestedDocument:
        from docx import Document
        doc = Document(str(path))

        md = _docx_to_markdown(doc)

        # Title: document core property → first heading → first paragraph
        title: str | None = None
        if doc.core_properties.title:
            title = doc.core_properties.title.strip() or None
        if not title:
            import re
            m = re.search(r"^#{1,6}\s+(.+)$", md, re.MULTILINE)
            if m:
                title = m.group(1).strip()
        if not title and doc.paragraphs:
            title = doc.paragraphs[0].text.strip()[:100] or None

        return IngestedDocument(
            text=md.strip(),
            title=title or path.stem,
            metadata={"author": doc.core_properties.author or ""},
        )
