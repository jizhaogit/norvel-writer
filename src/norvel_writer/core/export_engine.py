"""Exporters: NotebookLM, DOCX, Markdown."""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

log = logging.getLogger(__name__)


class BaseExporter:
    def export(
        self,
        project_id: str,
        dest_path: Path,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        raise NotImplementedError


class MarkdownExporter(BaseExporter):
    """Export all accepted chapter drafts as a single Markdown file."""

    def export(
        self,
        project_id: str,
        dest_path: Path,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        from norvel_writer.storage.db import get_db
        from norvel_writer.storage.repositories.project_repo import ProjectRepo
        from norvel_writer.storage.repositories.draft_repo import DraftRepo

        db = get_db()
        pr = ProjectRepo(db)
        dr = DraftRepo(db)

        project = pr.get_project(project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")

        chapters = pr.list_chapters(project_id)
        parts = [f"# {project['name']}\n"]

        for ch in chapters:
            parts.append(f"\n## {ch['title']}\n")
            draft = dr.get_accepted_draft(ch["id"])
            if draft:
                parts.append(draft["content"])
            else:
                parts.append("*(no draft)*")

        dest_path.parent.mkdir(parents=True, exist_ok=True)
        dest_path.write_text("\n\n".join(parts), encoding="utf-8")
        log.info("Exported Markdown to %s", dest_path)
        return dest_path


class DocxExporter(BaseExporter):
    """Export as a .docx file."""

    def export(
        self,
        project_id: str,
        dest_path: Path,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        from docx import Document as DocxDocument
        from norvel_writer.storage.db import get_db
        from norvel_writer.storage.repositories.project_repo import ProjectRepo
        from norvel_writer.storage.repositories.draft_repo import DraftRepo

        db = get_db()
        pr = ProjectRepo(db)
        dr = DraftRepo(db)

        project = pr.get_project(project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")

        doc = DocxDocument()
        doc.add_heading(project["name"], level=0)

        chapters = pr.list_chapters(project_id)
        for ch in chapters:
            doc.add_heading(ch["title"], level=1)
            draft = dr.get_accepted_draft(ch["id"])
            if draft:
                for para in draft["content"].split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph("(no draft)")

        dest_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest_path))
        log.info("Exported DOCX to %s", dest_path)
        return dest_path


class NotebookLMExporter(BaseExporter):
    """
    Export project sources in a NotebookLM-friendly format.
    Produces a structured Markdown file with sections and metadata.
    """

    def export(
        self,
        project_id: str,
        dest_path: Path,
        options: Optional[Dict[str, Any]] = None,
    ) -> Path:
        from norvel_writer.storage.db import get_db
        from norvel_writer.storage.repositories.project_repo import ProjectRepo
        from norvel_writer.storage.repositories.document_repo import DocumentRepo

        db = get_db()
        pr = ProjectRepo(db)
        doc_repo = DocumentRepo(db)

        project = pr.get_project(project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")

        opts = options or {}
        include_types = opts.get("doc_types", ["codex", "beats", "research", "notes"])

        parts = [
            f"# {project['name']} — Project Reference\n",
            f"*Language: {project.get('language', 'en')}*\n",
            f"*Description: {project.get('description', '')}*\n",
        ]

        for doc_type in include_types:
            docs = doc_repo.list_documents(project_id, doc_type=doc_type)
            if not docs:
                continue
            parts.append(f"\n## {doc_type.replace('_', ' ').title()}\n")
            for doc in docs:
                if doc.get("title"):
                    parts.append(f"\n### {doc['title']}\n")
                # Get first few chunks
                chunks = doc_repo.list_chunks(doc["id"])[:3]
                for chunk in chunks:
                    parts.append(chunk["text"])

        dest_path.parent.mkdir(parents=True, exist_ok=True)
        dest_path.write_text("\n\n".join(parts), encoding="utf-8")
        log.info("Exported NotebookLM package to %s", dest_path)
        return dest_path

    def open_in_browser(self) -> None:
        """Open NotebookLM in the default browser."""
        import webbrowser
        webbrowser.open("https://notebooklm.google.com")
